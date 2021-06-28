from tkinter import *
import requests
import docx
import xlrd2
from tkinter import messagebox
from datetime import datetime
import matplotlib.pyplot as plt
from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg

root = Tk() # Главное окно
root.title("Анализ уязвимостей ПО") # Название окна
root.geometry("1200x700") # Размер окна
root.configure(background='#0a3b4f') # Цвет заднего фона окна
labelDateInfo = Label(root, text="Анализ уязвимостей Adobe Photoshop", bg='#008080',
                      font='Times 20', fg='#613e3d', width=80)


def Download(event):
    operation_label.configure(state=NORMAL)
    operation_status_label['text'] = "Выполняется"
    operation_status_label.configure(state=NORMAL)
    root.update_idletasks()

    files = open('vullist.xlsx', "wb")

    url = 'https://bdu.fstec.ru/files/documents/vullist.xlsx'

    headers = {
        'User-Agent': 'My User Agent 1.0',
        'From': 'youremail@domain.com'  # This is another valid field
    }

    response = requests.get(url, headers=headers)
    files.write(response.content)
    files.close()
    operation_status_label.configure(text="Выполнено")
    messagebox.showinfo("Info", "Обновление базы данных выполнено")


def dateOn(event): # Функция для радиокнопки "По дате", включает поля для ввода даты.
    textBoxFromDate.configure(state=NORMAL)
    textBoxToDate.configure(state=NORMAL)
    labelFromDate.configure(state=NORMAL)
    labelToDate.configure(state=NORMAL)
    labelDate.configure(state=NORMAL)


def dateOff(event): # Функция для радиокнопки "Без даты", выключает и очищает поля для ввода даты.
    textBoxFromDate.delete(0, END)
    textBoxToDate.delete(0, END)
    textBoxFromDate.configure(state=DISABLED)
    textBoxToDate.configure(state=DISABLED)
    labelFromDate.configure(state=DISABLED)
    labelToDate.configure(state=DISABLED)
    labelDate.configure(state=DISABLED)


def AnalysisWithDate(event): # Функция для проверки правильности ввода даты
    radio_condition = radioButtonDateVar.get() # Заносим в переменную radio_condition состояние радиокнопок (1 или 0)

    if radio_condition == 1: # Если радиокнопка "По дате" включена (1)
        dataFrom = textBoxFromDate.get()
        dataTo = textBoxToDate.get()
        if len(dataFrom and dataTo) == 10 and (dataFrom[2] and dataTo[2]) == '.' and \
                (dataFrom[5] and dataTo[5]) == '.' and dataFrom[6:].isnumeric() and dataTo[6:].isnumeric() and \
                dataFrom[:2].isnumeric() and dataTo[:2].isnumeric() and dataFrom[3:5].isnumeric() and \
                dataTo[3:5].isnumeric():
            tsFrom = datetime(year=int(dataFrom[6:]), month=int(dataFrom[3:5]), day=int(dataFrom[:2]))
            tsTo = datetime(year=int(dataFrom[6:]), month=int(dataFrom[3:5]), day=int(dataFrom[:2]))
            if (tsFrom.date and tsTo.day) > 0 and (tsFrom.day and tsTo.day) < 32 and (tsFrom.month and tsTo.month) > 0 \
                    and (tsFrom.month and tsTo.month) < 13 and (tsFrom.year and tsTo.year) > 1900:
                Analysis(event)
            else:
                messagebox.showerror('Ошибка', 'Некорректно введена дата') # Если дата введена некорректно -
                                                                           # выводим окно с ошибкой
        else:
            messagebox.showerror('Ошибка', 'Некорректно введена дата')  # Если дата введена некорректно -
                                                                        # выводим окно с ошибкой
    else:
        Analysis(event) # Выполняем функцию Analysis


def Analysis(event): # Функция поиска уязвимостей
    analysis_status_label['state'] = NORMAL
    operation_analysis_status_label['text'] = "Выполняется"
    operation_analysis_status_label['state'] = NORMAL
    root.update_idletasks()
    workbook = xlrd2.open_workbook('vullist.xlsx')
    sheet = workbook.sheet_by_index(0)

    row = sheet.nrows  # определяем количество записей (строк) на листе
    if row == 0:
        messagebox.showerror("Error", "Необходимо обновить базу данных")
    else:
        print('Всего записей:', row)  # выведем количество записей на печать

        # выполним считывание списка данных из столбца с данными Название ПО
        names = sheet.col_values(4)  # (4-й столбец, нумерация с нуля)
        # выполним считывание списка данных из столбца с данными Уровень опасности
        danger_levels = sheet.col_values(12)  # (12-й столбец, нумерация с нуля)
        chrb = radioButtonDateVar.get()
        date_file = sheet.col_values(9)
        global name_software
        name_software = inputEntry.get()

        global danger_low, danger_middle, danger_hight, danger_super
        danger_super, danger_hight, danger_middle, danger_low = 0, 0, 0, 0  # инициализируем переменные-счетчики
                                                                            # различных уровней опасности
        if chrb == 0:  # Если радиокнопка По дате выключена (0)
            dataFrom = datetime.strptime('01.01.1900', '%d.%m.%Y')
            dataTo = datetime.strptime('17.06.3021', '%d.%m.%Y')
        else:
            dataFrom = datetime.strptime(textBoxFromDate.get(), '%d.%m.%Y')
            dataTo = datetime.strptime(textBoxToDate.get(), '%d.%m.%Y')

        for i in range(9, row):
            if date_file[i] != '':
                date_file[i] = datetime.strptime(date_file[i], '%d.%m.%Y')
            else:
                date_file[i] = datetime.strptime('01.01.1900', '%d.%m.%Y')

        for i in range(4, row):
            if (str(date_file[i]) >= str(dataFrom)) and (str(date_file[i]) <= str(dataTo)):
                if names[i].find(name_software) >= 0:      # если наименование ПО содержит искомое проверим по первой
                                                           # букве уровень уязвимости ПО
                    if danger_levels[i][0] == 'К':    # Критический
                        danger_super += 1
                    elif danger_levels[i][0] == 'В':  # Высокий
                        danger_hight += 1
                    elif danger_levels[i][0] == 'С':  # Средний
                        danger_middle += 1
                    else:                             # Низкий
                        danger_low += 1

        labelLowOut['text'] = danger_low
        labelMidOut['text'] = danger_middle
        labelHighOut['text'] = danger_hight
        labelSuperOut['text'] = danger_super
        operation_analysis_status_label['text'] = "Выполнено"
        messagebox.showinfo("Info ", "Анализ базы данных завершён")


def Clear(event): # Функция для очистки лейблов и полей
    labelLowOut['text'] = ""
    labelMidOut['text'] = ""
    labelHighOut['text'] = ""
    labelSuperOut['text'] = ""
    operation_status_label['text'] = ""
    operation_analysis_status_label['text'] = ""
    operation_status_label['state'] = DISABLED
    operation_label['state'] = DISABLED
    analysis_status_label['state'] = DISABLED
    textBoxFromDate.delete(0, END)
    textBoxToDate.delete(0, END)


def SaveDocx(event): # Функция для сохранения результатов в docx
    if (labelLowOut['text'] and labelMidOut['text'] and labelHighOut['text'] and labelSuperOut['text']) == "":
        messagebox.showerror("Error", "Сначала нужно провести анализ данных!")
    else:
        document = docx.Document()
        document.add_heading(name_software, 0)
        document.add_heading('Количество уязвимостей по уровням опасности', level=1)
        table = document.add_table(rows=4, cols=3)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '1'
        hdr_cells[1].text = 'Низкий'
        hdr_cells[2].text = str(labelLowOut['text'])
        hdr_cells1 = table.rows[1].cells
        hdr_cells1[0].text = '2'
        hdr_cells1[1].text = 'Средний'
        hdr_cells1[2].text = str(labelMidOut['text'])
        hdr_cells2 = table.rows[2].cells
        hdr_cells2[0].text = '3'
        hdr_cells2[1].text = 'Высокий'
        hdr_cells2[2].text = str(labelHighOut['text'])
        hdr_cells3 = table.rows[3].cells
        hdr_cells3[0].text = '4'
        hdr_cells3[1].text = 'Критический'
        hdr_cells3[2].text = str(labelSuperOut['text'])
        document.save('Анализ уязвимостей {}.docx'.format(name_software))


def SaveDocxEach(event):
    try:
        list_var = ["низкому", "среднему", "высокому", "критическому"]
        list_var_1 = ["низких", "средних", "высоких", "критических"]
        list_var_2 = ["Низкий", "Средний", "Высокий", "Критический"]

        for i in range(4):
            document = docx.Document()
            document.add_heading(name_software, 0)
            document.add_heading('Количество уязвимостей по {} уровню опасности '.format(list_var[i]), level=1)
            table = document.add_table(rows=1, cols=3)

            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = '1'
            hdr_cells[1].text = list_var_2[i]

            if i == 0:
                hdr_cells[2].text = str(labelLowOut['text'])
            elif i == 1:
                hdr_cells[2].text = str(labelMidOut['text'])
            elif i == 2:
                hdr_cells[2].text = str(labelHighOut['text'])
            else:
                hdr_cells[2].text = str(labelSuperOut['text'])

            document.save('Анализ {} уязвимостей {}.docx'.format(list_var_1[i], name_software))
    except NameError:
        messagebox.showerror("Error", "Сначала нужно провести анализ данных!")


def diagram(event):
    try:
        labels = ["Низкий", "Средний", "Высокий", "Критический"]
        sizes = [danger_low, danger_middle, danger_hight, danger_super]

        colors = ("grey", "yellow", "orange", "brown")
        fig1, ax1 = plt.subplots()
        explode = (0.05, 0.15, 0.1, 0.1)
        patches, texts, auto = ax1.pie(sizes, colors=colors, wedgeprops=dict(width=0.3),
                                       explode=explode, autopct='%1.1f%%')
        plt.legend(patches, labels, loc="best")
        window = Tk()
        window.title("Диаграмма уязвимостей")
        window.configure(background='#a8e4a0')
        canvas = FigureCanvasTkAgg(fig1, master=window)
        canvas.get_tk_widget().pack()
        canvas.draw()
    except NameError:
        messagebox.showerror("Error", "Для вывода диаграммы необходимо провести анализ!")


# Создание интерфейса
# Создание кнопок, полей, лейблов, их бинд и расположение
labelDateInfo.pack()

canvas_settings = Canvas(root, width=558, height=227, background="#732626")
canvas_settings.place(x=0, y=38)

canvas_analysis_result = Canvas(root, width=558, height=170, background="#722626")
canvas_analysis_result.place(x=0, y=298)

canvas_function = Canvas(root, width=266, height=227, background="#722626")
canvas_function.place(x=598, y=38)

canvas_save = Canvas(root, width=276, height=227, background="#722626")
canvas_save.place(x=898, y=38)

radioButtonDateVar = BooleanVar() # Создание радиокнопок
radioButtonDateVar.set(0)

radioButtonDateOn = Radiobutton(root, text="По дате", bg='#447185', variable=radioButtonDateVar, value=1)
radioButtonDateOn.bind('<Button-1>', dateOn) #Привязка функции "dateOn" к радиокнопке "По дате"

radioButtonDateOff = Radiobutton(root, text="За все время", bg='#447185', variable=radioButtonDateVar, value=0)
radioButtonDateOff.bind('<Button-1>', dateOff)

radioButtonDateOff.place(x=450, y=75)
radioButtonDateOn.place(x=45, y=75)

analysis_result = Label(root, background='#e6dd43', foreground='#29281d',
                        text="Результаты анализа данных:", font=('Times', 16), width=46)
analysis_result.place(x=2, y=300)

function_label = Label(root, text="Функции", background='orange', font=('Times', 14), width=26)
function_label.place(x=600, y=40)

buttonAnalysis = Button(root, background='#709972', font='Times 12', text="Анализ", width=13, height=2)
buttonAnalysis.bind('<Button-1>', AnalysisWithDate) #Привязка функции "AnalysisWithDate" к кнопке "Анализ"
buttonAnalysis.place(x=601, y=140) #Размещаем кнопки по координатам на плоскости окна

buttonobnow = Button(root, bg='#ffd35f', font='Times 12', text="Обновить базу", width=13, height=2)
buttonobnow.bind('<Button-1>', Download)
buttonobnow.place(x=601, y=70)

buttonClear = Button(root, bg='#008B8B', font='Times 12', text="Очистить поля", width=13, height=2)
buttonClear.bind('<Button-1>', Clear) #Привязка функции "Clear" к кнопке "Очистить все"
buttonClear.place(x=1000, y=600) #Размещаем кнопки по координатам на плоскости окна

buttonSave = Button(root, background='#e6a87c', font='Times 12', text="Сохранить всё \n в один", width=13, height=2)
buttonSave.bind('<Button-1>', SaveDocx)  #Привязка функции "SaveDocx" к кнопке "Сохранить в docx"
buttonSave.place(x=901, y=70) #Размещаем кнопки по координатам на плоскости окна

buttonSaveEach = Button(root, background='#e6a87c', font='Times 12', text="Сохранить каждый \n отдельно",
                        width=14, height=2)
buttonSaveEach.bind('<Button-1>', SaveDocxEach)
buttonSaveEach.place(x=1040, y=70) #Размещаем кнопки по координатам на плоскости окна

labelLow = Label(root, width=13, height=2, bg='#00ff00', font='Times 13', text="Низкий")
labelLow.place(x=15, y=350)

labelLowOut = Label(root, bg='#ffffff', font='Times 15', fg='black', width=5)
labelLowOut.place(x=48, y=405)

labelMid = Label(root, width=13, height=2, bg='#0099ff', font='Times 13', text="Средний")
labelMid.place(x=150, y=350)

labelMidOut = Label(root, bg='#ffffff', font='Times 15', fg='black', width=5)
labelMidOut.place(x=181, y=405)

labelHigh = Label(root, width=13, height=2, bg='#ff0099', font='Times 13', text="Высокий")
labelHigh.place(x=285, y=350)

labelHighOut = Label(root, bg='#ffffff', font='Times 15', fg='black', width=5)
labelHighOut.place(x=317, y=405)

labelSuper = Label(root, width=13, height=2, bg='#ff0000', font='Times 13', text="Критический")
labelSuper.place(x=420, y=350)

labelSuperOut = Label(root, bg='#ffffff', font='Times 15', fg='black', width=5)
labelSuperOut.place(x=451, y=405)

labelDate = Label(root, text="Введите необходимую дату:", state=DISABLED,
                  bg='#FFFAFA', font='Times 13', fg='#000', width=30)
labelDate.place(x=145, y=75)

labelFromDate = Label(root, text=" От:", state=DISABLED, bg='#FFFAFA', fg='black', width=5)
labelFromDate.place(x=155, y=115)

labelToDate = Label(root, text="До:", state=DISABLED, bg='#FFFAFA', fg='black', width=5)
labelToDate.place(x=295, y=115)

textBoxFromDate = Entry(root, state=DISABLED, width=10)
textBoxFromDate.place(x=205, y=115, height=21)

textBoxToDate = Entry(root, state=DISABLED, width=10)
textBoxToDate.place(x=345, y=115, height=21)

button_diagram = Button(root, bg='#7fc7ff', font='Times 12', text="Вывести \n диаграмму", height=2, width=13)
button_diagram.bind('<Button-1>', diagram)
button_diagram.place(x=601, y=215)

inputLabel = Label(root, background="#1689e0", font='Times 11', text="Название ПО", width=13)
inputLabel.place(x=225, y=165)

inputEntry = Entry(root, width=21)
inputEntry.insert(0, "Adobe Photoshop")
inputEntry.place(x=212, y=195)

operation_label = Label(root, background="#d9a925", state=DISABLED, text="Статус обновления:", width=16)
operation_label.place(x=740, y=73)

operation_status_label = Label(root, background="#97fa02", state=DISABLED, text="", width=16)
operation_status_label.place(x=740, y=97)

analysis_status_label = Label(root, background="#d9a925", state=DISABLED, text="Статус анализа:", width=16)
analysis_status_label.place(x=740, y=143)

operation_analysis_status_label = Label(root, background="#97fa02", state=DISABLED, text="", width=16)
operation_analysis_status_label.place(x=740, y=167)

save_label = Label(root, text="Варианты сохранения в файл", font='Times 13', background="orange", width=30)
save_label.place(x=900, y=40)

settings_label = Label(root, text="Настройки анализа данных", background="#9e9b9b", foreground="#3b00ff",
                       font=('Times', 16), width=46)
settings_label.place(x=2, y=40)

print("Круг программы выполнен")
root.mainloop()
